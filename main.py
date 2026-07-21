import os
import shutil
import zipfile
import io
import json
import time
import asyncio
import hashlib
import logging
from pathlib import Path
from datetime import datetime, timedelta
from dotenv import load_dotenv
from docx import Document
import requests
import httpx
import pdfplumber
from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request, Response
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from urllib.parse import quote
import openpyxl
import xlrd
import uvicorn
import database
import parser
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, validator
from typing import List, Optional
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ================= НОВАЯ БИБЛИОТЕКА ЮKASSA =================
from yookassa import Configuration, Payment

# ================= ЗАГРУЗКА .ENV =================
load_dotenv()

# ================= ЛОГИРОВАНИЕ =================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ================= НАСТРОЙКИ =================
API_KEY = os.getenv("DEEPSEEK_API_KEY")
if not API_KEY:
    raise ValueError("DEEPSEEK_API_KEY не найден в .env")

OLLAMA_API_URL = "https://api.deepseek.com/v1/chat/completions"
MODEL_NAME = "deepseek-chat"
MAX_TENDERS = 15
ADMIN_TOKEN = os.getenv("ADMIN_TOKEN")
if not ADMIN_TOKEN:
    raise ValueError("ADMIN_TOKEN не найден в .env")

# ================= НАСТРОЙКИ ЮKASSA =================
YOOKASSA_SHOP_ID = os.getenv("YOOKASSA_SHOP_ID")
YOOKASSA_SECRET_KEY = os.getenv("YOOKASSA_SECRET_KEY")
if not all([YOOKASSA_SHOP_ID, YOOKASSA_SECRET_KEY]):
    raise ValueError("ЮKassa настройки не найдены в .env")

Configuration.account_id = YOOKASSA_SHOP_ID
Configuration.secret_key = YOOKASSA_SECRET_KEY

# ================= НАСТРОЙКИ MAX BOT =================
MAX_BOT_TOKEN = os.getenv("MAX_BOT_TOKEN")
MAX_CHAT_ID = os.getenv("MAX_CHAT_ID")
MAX_API_URL = "https://platform-api2.max.ru/messages"
if not all([MAX_BOT_TOKEN, MAX_CHAT_ID]):
    raise ValueError("MAX BOT настройки не найдены в .env")

# ================= FASTAPI + RATE LIMITING =================
limiter = Limiter(key_func=get_remote_address)
app = FastAPI()
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ================= CORS =================
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ================= MIDDLEWARE ЛОГИРОВАНИЯ =================
@app.middleware("http")
async def log_requests(request: Request, call_next):
    start = time.time()
    response = await call_next(request)
    duration = time.time() - start
    license_header = request.headers.get("X-License-Key", "none")
    logger.info(
        f"{request.method} {request.url.path} - {response.status_code} - "
        f"{duration:.2f}s - {request.client.host} - "
        f"License: {license_header[:8]}..."
    )
    return response

# ================= PYDANTIC МОДЕЛИ =================
class TenderData(BaseModel):
    url: str = Field(..., max_length=500)
    regNumber: str = Field(..., max_length=50)
    text: str = Field(..., max_length=50000)

    @validator("url")
    def validate_url(cls, v):
        if not v.startswith(("http://", "https://")):
            raise ValueError("URL должен начинаться с http:// или https://")
        return v

class AnalyzeRequest(BaseModel):
    tenders: List[TenderData] = Field(..., max_items=50)
    fields: List[str] = Field(default_factory=list, max_items=20)

class TrialRequest(BaseModel):
    device_id: str = Field(..., min_length=10, max_length=100)

class LicenseActivateRequest(BaseModel):
    key: str = Field(..., min_length=10, max_length=100)

# ================= ИНИЦИАЛИЗАЦИЯ БД =================
@app.on_event("startup")
async def startup():
    try:
        await database.init_db()
        logger.info("База данных PostgreSQL инициализирована")
    except Exception as e:
        logger.error(f"Ошибка инициализации БД: {e}")
        raise

# ================= ПРОВЕРКА ДОСТУПА =================
async def check_access(request: Request):
    license_key = request.headers.get("X-License-Key")
    if license_key:
        result = await database.verify_license(license_key)
        if result and result.get("valid"):
            return {"type": "license", "id": license_key}
    device_id = request.headers.get("X-Device-ID")
    if device_id:
        is_active = await database.check_trial_by_device(device_id)
        if is_active:
            return {"type": "trial", "id": device_id}
    logger.warning(f"Неавторизованный доступ с {request.client.host}")
    raise HTTPException(401, detail="Требуется действующая лицензия или активный пробный период")

# ================= ФУНКЦИЯ ОТПРАВКИ В MAX =================
async def send_max_notification(
    reg_number: str,
    client_name: str,
    inn: str,
    phone: str,
    email: str,
    nmc: str,
    end_date: str,
    bid_end_date: str,
    bid_security: str,
    contract_security: str,
    guarantee_type: str,
    contact_by_email: bool
):
    if not MAX_BOT_TOKEN or not MAX_CHAT_ID:
        logger.warning("MAX Bot не настроен: пропуск уведомления")
        return

    text = (
        f"Новая заявка на банковскую гарантию!\n\n"
        f"Номер тендера: {reg_number}\n"
        f"Клиент: {client_name}\n"
        f"ИНН: {inn or 'Не указан'}\n"
        f"Телефон: {phone}\n"
        f"Email: {email}\n"
        f"Начальная цена (НМЦ): {nmc or 'Не указана'}\n"
        f"Дата окончания контракта: {end_date or 'Не указана'}\n"
        f"Дата окончания подачи заявок: {bid_end_date or 'Не указана'}\n"
        f"Обеспечение заявки: {bid_security or 'Не указано'}\n"
        f"Обеспечение контракта: {contract_security or 'Не указано'}\n"
        f"Тип гарантии: {'Обеспечение заявки (участие)' if guarantee_type == 'participation' else 'Обеспечение исполнения контракта'}\n"
        f"Связь только по email: {'Да' if contact_by_email else 'Нет (звонить)'}\n"
        f"Время заявки: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )

    url = f"{MAX_API_URL}?chat_id={MAX_CHAT_ID}"
    headers = {
        "Authorization": MAX_BOT_TOKEN,
        "Content-Type": "application/json"
    }
    payload = {"text": text}

    try:
        async with httpx.AsyncClient(timeout=10, verify=False) as client:
            response = await client.post(url, headers=headers, json=payload)
            if response.status_code == 200:
                logger.info(f"Уведомление в MAX отправлено (тендер {reg_number})")
            else:
                logger.error(f"Ошибка MAX API: {response.status_code} - {response.text}")
    except Exception as e:
        logger.error(f"Исключение при отправке в MAX: {e}")

# ================= ФУНКЦИИ ЧТЕНИЯ ФАЙЛОВ =================
def read_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"Ошибка чтения docx: {e}"

def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        try:
            with open(file_path, 'r', encoding='cp1251') as f:
                return f.read()
        except Exception as e:
            return f"Ошибка чтения txt: {e}"

def read_excel(file_path):
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheet = wb.active
        rows = sheet.iter_rows(values_only=True)
        return "\n".join(str(cell) for row in rows for cell in row if cell)
    except:
        try:
            wb = xlrd.open_workbook(file_path)
            sheet = wb.sheet_by_index(0)
            text = ""
            for row in range(sheet.nrows):
                row_text = [str(sheet.cell_value(row, col)) for col in range(sheet.ncols) if sheet.cell_value(row, col)]
                if row_text:
                    text += " ".join(row_text) + "\n"
            return text
        except Exception as e:
            return f"Ошибка чтения Excel: {e}"

def read_pdf(file_path):
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        text = text.strip()
        if not text:
            logger.info(f"⚠️ [PDF] {file_path} не содержит текста (возможно, сканированный)")
        return text if text else "PDF не содержит текста (возможно, сканированный документ)"
    except Exception as e:
        logger.error(f"❌ [PDF] Ошибка при открытии {file_path}: {e}")
        return f"Ошибка чтения PDF: {e}"

# ================= ОПРЕДЕЛЕНИЕ ТИПА ФАЙЛА =================
def detect_file_type(content: bytes, filename: str = "") -> str:
    ext = os.path.splitext(filename)[1].lower() if filename else ""

    if content.startswith(b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'):
        if b'WorkBook' in content[:2000] or b'BOUNDSHEET' in content[:2000]:
            return 'xls'
        return 'doc'

    if ext == '.doc':
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                if any(f.startswith('word/') for f in zf.namelist()):
                    return 'docx'
        except:
            pass
        return 'doc'

    if content.startswith(b'%PDF'):
        return 'pdf'

    if content.startswith(b'PK\x03\x04') or content.startswith(b'PK\x05\x06') or content.startswith(b'PK\x07\x08'):
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                files = zf.namelist()
                if any(f.startswith('word/') for f in files):
                    return 'docx'
                if any(f.startswith('xl/') for f in files):
                    return 'xlsx'
                return 'zip'
        except:
            return 'zip'

    if content.startswith(b'Rar!\x1a\x07\x00') or content.startswith(b'Rar!\x1a\x07\x01\x00') or content.startswith(b'Rar!'):
        return 'rar'
    if content.startswith(b'7z\xbc\xaf\x27\x1c'):
        return '7z'
    if content.startswith(b'\x89PNG\r\n\x1a\n'):
        return 'png'
    if content.startswith(b'\xff\xd8\xff'):
        return 'jpg'
    if content.startswith(b'{\\rtf'):
        return 'rtf'

    if ext == '.doc':
        return 'doc'
    return 'unknown'

# ================= ЗАПРОСЫ К DEEPSEEK =================
def query_deepseek(prompt, license_key=None, device_id=None):
    messages = [{"role": "system", "content": "Ты — эксперт по анализу тендерной документации. Отвечай чётко, по делу, без воды."},
                {"role": "user", "content": prompt}]
    payload = {"model": MODEL_NAME, "messages": messages, "temperature": 0.3, "max_tokens": 2500, "stream": False}
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    try:
        response = requests.post(OLLAMA_API_URL, json=payload, headers=headers, timeout=180)
        total_tokens = 0
        if response.status_code == 200:
            result = response.json()
            usage = result.get("usage", {})
            total_tokens = usage.get("total_tokens", 0)
            answer = result["choices"][0]["message"]["content"]
        else:
            answer = f"Ошибка HTTP {response.status_code}: {response.text}"
        if license_key or device_id:
            asyncio.create_task(database.increment_usage(license_key=license_key, device_id=device_id, tokens_used=total_tokens))
        return answer
    except Exception as e:
        error_msg = f"Ошибка: {e}"
        if license_key or device_id:
            asyncio.create_task(database.increment_usage(license_key=license_key, device_id=device_id, tokens_used=0))
        return error_msg

def analyze_tender_text(text, selected_fields, license_key=None, device_id=None, max_text_len=8000):
    if not selected_fields:
        selected_fields = [
            "НАЗВАНИЕ АУКЦИОНА", "Начальная цена (НМЦ)", "ДОПОЛНИТЕЛЬНЫЕ ТРЕБОВАНИЯ К УЧАСТНИКУ",
            "ДАТА ОКОНЧАНИЯ/ПРОВЕДЕНИЯ", "Аванс", "Обеспечение заявки", "Обеспечение контракта",
            "Обеспечение гарантийных обязательств", "Контакты", "Место исполнения", "ДАТА ОКОНЧАНИЯ КОНТРАКТА"
        ]
    fields_str = "\n".join(f"{field}: " for field in selected_fields)
    truncated = text[:max_text_len]
    prompt = f"""Ты анализируешь тендерную документацию. Извлеки из текста следующие данные. Если информации нет, напиши "Информация отсутствует".
Ответ должен быть строго в таком формате (каждый пункт с новой строки):
{fields_str}
Вот текст для анализа:
{truncated}
Извлеки данные и напиши в указанном формате."""
    answer = query_deepseek(prompt, license_key=license_key, device_id=device_id)
    result = {}
    for line in answer.split('\n'):
        if ':' in line:
            k, v = line.split(':', 1)
            result[k.strip()] = v.strip()
    return result

def critical_analysis(text, license_key=None, device_id=None, custom_prompt=None):
    if custom_prompt and custom_prompt.strip():
        prompt = custom_prompt.strip()
        prompt += f"\n\nТекст тендера и прикреплённых документов:\n{text}"
    else:
        prompt = f"""Ты — эксперт по тендерной документации. Проанализируй текст тендера и выяви потенциальные риски, сложности, скрытые требования, которые могут помешать выполнению контракта. Оцени, насколько выполним данный тендер для среднестатистического поставщика.

    Текст тендера и документов:
    {text}

    Ответ должен быть в виде связного текста (5–7 предложений), где будут перечислены:
    - основные риски и сложности,
    - неочевидные требования,
    - рекомендации по успешному выполнению.
    Не используй маркированные списки, пиши сплошным текстом.
    """
    answer = query_deepseek(prompt, license_key=license_key, device_id=device_id)
    return answer

# ================= ЭНДПОЙНТЫ HEALTH CHECK =================
@app.get("/health")
async def health():
    return {"status": "ok"}

# ================= СТРАНИЦЫ САЙТА =================
@app.get("/", response_class=HTMLResponse)
async def main_page():
    with open("templates/main.html", "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/contacts", response_class=HTMLResponse)
async def contacts_page():
    with open("templates/contacts.html", "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/offer", response_class=HTMLResponse)
async def offer_page():
    with open("templates/offer.html", "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/oferta")
async def download_oferta():
    file_path = os.path.join("static", "oferta.docx")
    if not os.path.exists(file_path):
        raise HTTPException(404, "Файл оферты не найден")
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="oferta.docx")

@app.get("/privacy-policy", response_class=HTMLResponse)
async def privacy_policy_page():
    with open("templates/privacy-policy.html", "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/download-file/{filename}")
async def download_file(filename: str):
    allowed_files = ["tender-parser-extension.crx", "tender-parser-extension.zip"]
    if filename not in allowed_files:
        raise HTTPException(404, "Файл не найден")
    file_path = os.path.join(os.path.dirname(__file__), filename)
    if not os.path.exists(file_path):
        raise HTTPException(404, "Файл не найден на сервере")
    media_type = "application/x-chrome-extension" if filename.endswith(".crx") else "application/zip"
    return FileResponse(file_path, media_type=media_type, filename=filename)

@app.get("/download")
async def download_redirect():
    return RedirectResponse(url="/", status_code=302)

@app.get("/updates.xml")
async def updates_xml():
    file_path = os.path.join(os.path.dirname(__file__), "updates.xml")
    if not os.path.exists(file_path):
        raise HTTPException(404, "Файл updates.xml не найден")
    return FileResponse(file_path, media_type="application/xml")

# ================= ЭНДПОЙНТЫ ЮKASSA =================
@app.post("/api/create-payment")
@limiter.limit("5/minute")
async def create_payment(request: Request, data: dict = None):
    amount_value = data.get("amount", 2500) if data else 2500
    device_id = request.headers.get("X-Device-ID", "unknown")
    email = data.get("email", "client@example.com") if data else "client@example.com"
    try:
        payment = Payment.create({
            "amount": {"value": str(amount_value), "currency": "RUB"},
            "confirmation": {
                "type": "redirect",
                "return_url": f"https://csb24-tender.ru/success?device_id={device_id}",
                "cancel_url": "https://csb24-tender.ru/"
            },
            "capture": True,
            "description": "Лицензия для Тендерного парсера (1 месяц)",
            "metadata": {"device_id": device_id},
            "receipt": {
                "customer": {"email": email},
                "items": [{
                    "description": "Лицензия на использование Тендерного парсера (1 месяц)",
                    "quantity": "1.00",
                    "amount": {"value": str(amount_value), "currency": "RUB"},
                    "vat_code": 4,
                    "payment_mode": "full_payment",
                    "payment_subject": "service"
                }]
            }
        })
        await database.save_payment(payment.id, device_id)
        logger.info(f"Платёж создан: id={payment.id}, сумма={payment.amount.value} RUB")
        return {"id": payment.id, "confirmation_url": payment.confirmation.confirmation_url, "amount": payment.amount.value}
    except Exception as e:
        logger.error(f"Ошибка создания платежа: {e}")
        raise HTTPException(500, f"Не удалось создать платёж: {str(e)}")

@app.post("/yookassa-webhook")
@limiter.limit("10/minute")
async def yookassa_webhook(request: Request):
    body = await request.json()
    event = body.get("event")
    if not event:
        raise HTTPException(400, "Неверный запрос")
    logger.info(f"Получен вебхук: {event}")
    if event == "payment.succeeded":
        payment_obj = body["object"]
        payment_id = payment_obj["id"]
        metadata = payment_obj.get("metadata", {})
        device_id = metadata.get("device_id", "unknown")
        license_key = await database.create_license(email=device_id, days_valid=30)
        if license_key:
            await database.update_payment_license(payment_id, license_key)
            logger.info(f"Лицензия {license_key} создана для device_id {device_id}")
        else:
            logger.warning(f"Не удалось создать лицензию для {device_id}")
        return {"status": "ok"}
    return {"status": "ok"}

@app.get("/success", response_class=HTMLResponse)
async def success_page(request: Request):
    device_id = request.query_params.get("device_id")
    if not device_id:
        return HTMLResponse("<html><body><h1>Ошибка</h1><p>Не указан идентификатор устройства.</p></body></html>")
    if device_id == "unknown":
        return HTMLResponse("<html><body><h1>Ошибка</h1><p>Идентификатор устройства не определён.</p></body></html>")
    license_key = await database.get_license_by_device(device_id)
    if not license_key:
        return HTMLResponse(f"""
        <!DOCTYPE html>
        <html>
        <head><meta charset="UTF-8"><title>Оплата успешна</title>
        <style>body{{font-family:Arial;text-align:center;padding:40px}}.key{{background:#f1f5f9;padding:15px;border-radius:8px;font-size:20px;font-weight:bold;letter-spacing:2px}}</style>
        </head>
        <body>
            <h1 style="color:#10b981;">✅ Оплата прошла успешно!</h1>
            <p>Ваш лицензионный ключ:</p>
            <div class="key" id="licenseKey">Загрузка...</div>
            <p><a href="/">На главную</a></p>
            <script>
                (async function(){{
                    const deviceId="{device_id}";
                    let attempts=0;
                    while(attempts<15){{
                        try{{
                            const resp=await fetch(`/api/get-license-by-device?device_id=${{deviceId}}`);
                            const data=await resp.json();
                            if(data.license_key){{
                                document.getElementById('licenseKey').textContent=data.license_key;
                                return;
                            }}
                        }}catch(e){{}}
                        attempts++;
                        await new Promise(r=>setTimeout(r,2000));
                    }}
                    document.getElementById('licenseKey').textContent='Не удалось получить ключ. Обратитесь в поддержку.';
                }})();
            </script>
        </body>
        </html>
        """)
    return HTMLResponse(f"""
    <!DOCTYPE html>
    <html>
    <head><meta charset="UTF-8"><title>Оплата успешна</title>
    <style>body{{font-family:Arial;text-align:center;padding:40px}}.key{{background:#f1f5f9;padding:15px;border-radius:8px;font-size:20px;font-weight:bold;letter-spacing:2px}}</style>
    </head>
    <body>
        <h1 style="color:#10b981;">✅ Оплата прошла успешно!</h1>
        <p>Ваш лицензионный ключ:</p>
        <div class="key">{license_key}</div>
        <p><a href="/">На главную</a></p>
    </body>
    </html>
    """)

@app.get("/api/get-license-by-device")
@limiter.limit("10/minute")
async def get_license_by_device(request: Request, device_id: str):
    if not device_id:
        raise HTTPException(400, "Не указан device_id")
    license_key = await database.get_license_by_device(device_id)
    if not license_key:
        raise HTTPException(404, "Лицензия ещё не создана")
    return {"license_key": license_key}

# ================= ЛИЦЕНЗИИ =================
@app.post("/api/create-order")
@limiter.limit("10/minute")
async def create_order(request: Request):
    admin_token = request.headers.get("X-Admin-Token")
    if admin_token != ADMIN_TOKEN:
        raise HTTPException(403, "Неверный admin token")
    license_key = await database.create_license(days_valid=30)
    if not license_key:
        raise HTTPException(500, "Не удалось создать лицензию")
    result = await database.verify_license(license_key)
    expires_at = result.get("expires_at") if result and result.get("valid") else None
    return {"status": "success", "license_key": license_key, "expires_at": expires_at}

@app.post("/api/admin/create-license")
@limiter.limit("5/minute")
async def admin_create_license(request: Request, data: dict = None):
    admin_token = request.headers.get("X-Admin-Token")
    if admin_token != ADMIN_TOKEN:
        raise HTTPException(403, "Invalid admin token")
    days = data.get("days", 30) if data else 30
    license_key = await database.create_license(days_valid=days)
    if not license_key:
        raise HTTPException(500, "Не удалось создать лицензию")
    return {"license_key": license_key, "days": days, "expires_at": (datetime.now() + timedelta(days=days)).isoformat()}

app.mount("/static", StaticFiles(directory="static"), name="static")

# ================= ЭНДПОЙНТ АНАЛИЗА ТЕКСТОВ (МАССОВЫЙ) =================
@app.post("/analyze_texts")
@limiter.limit("10/minute")
async def analyze_texts(request: Request, data: AnalyzeRequest):
    await check_access(request)
    tenders_data = data.tenders
    if not tenders_data:
        raise HTTPException(400, "Нет данных для анализа")
    selected_fields = data.fields
    if not selected_fields:
        selected_fields = [
            "НАЗВАНИЕ АУКЦИОНА", "Начальная цена (НМЦ)", "ДОПОЛНИТЕЛЬНЫЕ ТРЕБОВАНИЯ К УЧАСТНИКУ",
            "ДАТА ОКОНЧАНИЯ/ПРОВЕДЕНИЯ", "Аванс", "Обеспечение заявки", "Обеспечение контракта",
            "Обеспечение гарантийных обязательств", "Контакты", "Место исполнения", "ДАТА ОКОНЧАНИЯ КОНТРАКТА"
        ]
    license_key = request.headers.get("X-License-Key")
    device_id = request.headers.get("X-Device-ID")

    async def analyze_one(tender):
        reg_number = tender.regNumber
        cached = await database.get_cached_analysis(reg_number)
        if cached:
            await database.increment_cache_usage(reg_number)
            logger.info(f"Кэш для {reg_number} использован")
            return {"url": tender.url, "reg_number": reg_number, "analysis": cached}
        tender_text = tender.text
        if not tender_text or len(tender_text) < 100:
            return {"url": tender.url, "reg_number": reg_number, "error": "Недостаточно текста для анализа"}
        start = time.time()
        analysis_result = analyze_tender_text(tender_text, selected_fields, license_key, device_id)
        await database.save_analysis_cache(reg_number, analysis_result)
        logger.info(f"DeepSeek обработал {reg_number} за {time.time()-start:.2f} сек")
        return {"url": tender.url, "reg_number": reg_number, "analysis": analysis_result}

    tasks = [analyze_one(t) for t in tenders_data]
    results = await asyncio.gather(*tasks)

    doc = Document()
    doc.add_heading('РЕЗУЛЬТАТЫ АНАЛИЗА ТЕНДЕРОВ', 0)
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
    doc.add_paragraph('=' * 50)
    for item in results:
        doc.add_heading(f'Тендер: {item.get("reg_number", "Неизвестно")}', level=1)
        if "error" in item:
            doc.add_paragraph(f'Ошибка: {item["error"]}')
        else:
            analysis = item.get("analysis", {})
            if isinstance(analysis, dict):
                for k, v in analysis.items():
                    doc.add_paragraph(f'{k}: {v}')
            else:
                doc.add_paragraph(str(analysis))
        doc.add_page_break()

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zf:
        zf.writestr(f'результаты_анализа_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx', word_buffer.getvalue())
    zip_buffer.seek(0)
    filename = f'результаты_анализа_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip'
    encoded = quote(filename)
    return Response(zip_buffer.getvalue(), media_type="application/zip", headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"})

# ================= ЭНДПОЙНТ ДЛЯ АНАЛИЗА ТЕКУЩЕГО ТЕНДЕРА С ДОКУМЕНТАМИ (С ТАЙМАУТОМ НА ФАЙЛЫ) =================
@app.post("/analyze_tender_with_files")
@limiter.limit("10/minute")
async def analyze_tender_with_files(
    request: Request,
    regNumber: str = Form(...),
    printFormText: str = Form(...),
    fields: str = Form(...),
    files: list[UploadFile] = File(...),
    custom_critical_prompt: str = Form("")
):
    FILE_READ_TIMEOUT = 13  # секунд на чтение одного файла

    start_total = time.time()
    logger.info(f"🚀 [START] Тендер {regNumber}, файлов: {len(files)}")

    try:
        await check_access(request)
    except Exception as e:
        logger.error(f"❌ [ACCESS] Ошибка доступа для {regNumber}: {e}")
        raise

    selected_fields = json.loads(fields)
    if not selected_fields:
        selected_fields = [
            "НАЗВАНИЕ АУКЦИОНА", "Начальная цена (НМЦ)", "ДОПОЛНИТЕЛЬНЫЕ ТРЕБОВАНИЯ К УЧАСТНИКУ",
            "ДАТА ОКОНЧАНИЯ/ПРОВЕДЕНИЯ", "Аванс", "Обеспечение заявки", "Обеспечение контракта",
            "Обеспечение гарантийных обязательств", "Контакты", "Место исполнения", "ДАТА ОКОНЧАНИЯ КОНТРАКТА"
        ]

    device_id = request.headers.get("X-Device-ID", "unknown")
    license_key = request.headers.get("X-License-Key", None)

    prompt_info = "стандартного"
    if custom_critical_prompt and custom_critical_prompt.strip():
        prompt_info = "вашего личного"

    combined_text = printFormText
    original_files = []

    temp_dir = Path(f"/tmp/tender_{regNumber}_{datetime.now().strftime('%Y%m%d%H%M%S')}")
    temp_dir.mkdir(parents=True, exist_ok=True)

    try:
        start_files = time.time()
        for idx, file in enumerate(files):
            logger.info(f"📄 [FILE] Обработка файла {idx+1}/{len(files)}: {file.filename}")
            content = await file.read()

            file_type = detect_file_type(content, file.filename)
            ext_map = {
                'pdf': '.pdf', 'xlsx': '.xlsx', 'xls': '.xls', 'docx': '.docx',
                'rar': '.rar', 'zip': '.zip', '7z': '.7z', 'png': '.png',
                'jpg': '.jpg', 'rtf': '.rtf', 'doc': '.doc'
            }
            new_ext = ext_map.get(file_type, '')
            base, old_ext = os.path.splitext(file.filename)
            corrected_filename = base + new_ext if new_ext else file.filename

            file_path = temp_dir / corrected_filename
            with open(file_path, "wb") as f:
                f.write(content)

            ext = file_path.suffix.lower()
            text = ""
            try:
                if ext == ".docx":
                    text = await asyncio.wait_for(
                        asyncio.to_thread(read_docx, str(file_path)),
                        timeout=FILE_READ_TIMEOUT
                    )
                    logger.info(f"📄 [DOCX] {file.filename} -> {len(text)} символов извлечено")
                elif ext == ".txt":
                    text = await asyncio.wait_for(
                        asyncio.to_thread(read_txt, str(file_path)),
                        timeout=FILE_READ_TIMEOUT                    )
                    if text:
                        logger.info(f"📄 [TXT] {file.filename} -> {len(text)} символов извлечено")
                elif ext in (".xlsx", ".xls"):
                    text = await asyncio.wait_for(
                        asyncio.to_thread(read_excel, str(file_path)),
                        timeout=FILE_READ_TIMEOUT
                    )
                    logger.info(f"📄 [EXCEL] {file.filename} -> {len(text)} символов извлечено")
                elif ext == ".pdf":
                    text = await asyncio.wait_for(
                        asyncio.to_thread(read_pdf, str(file_path)),
                        timeout=FILE_READ_TIMEOUT
                    )
                    logger.info(f"📄 [PDF] {file.filename} -> {len(text)} символов извлечено")
                else:
                    text = ""
            except asyncio.TimeoutError:
                logger.error(f"⏱️ [TIMEOUT] Чтение файла {file.filename} превысило {FILE_READ_TIMEOUT} сек, пропускаем")
                text = f"[Пропущено: таймаут чтения {FILE_READ_TIMEOUT} сек]"
            except Exception as e:
                logger.error(f"❌ [ERROR] Ошибка чтения {file.filename}: {e}")
                text = f"[Ошибка чтения: {e}]"

            if text and not text.startswith("Ошибка") and not text.startswith("[Пропущено") and not text.startswith("[Ошибка"):
                combined_text += f"\n\n--- Содержимое файла {corrected_filename} ---\n{text}"

            original_files.append((corrected_filename, content))

        logger.info(f"⏱ [FILES] Обработка файлов завершена за {time.time() - start_files:.2f} сек")

    except Exception as e:
        logger.error(f"❌ [FILES] Критическая ошибка при обработке файлов: {e}")
        raise
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

    start_analysis = time.time()
    logger.info(f"🧠 [AI] Начало анализа текста для {regNumber}, длина: {len(combined_text)} символов")
    try:
        analysis_result = analyze_tender_text(combined_text, selected_fields, license_key, device_id, max_text_len=1000000)
    except Exception as e:
        logger.error(f"❌ [AI] Ошибка при структурированном анализе: {e}")
        raise

    try:
        critical_result = critical_analysis(combined_text, license_key, device_id, custom_prompt=custom_critical_prompt)
    except Exception as e:
        logger.error(f"❌ [AI] Ошибка при критическом анализе: {e}")
        raise

    logger.info(f"⏱ [AI] Анализ завершён за {time.time() - start_analysis:.2f} сек")

    await database.save_analysis_cache(regNumber, analysis_result)

    doc = Document()
    doc.add_heading('РЕЗУЛЬТАТЫ АНАЛИЗА ТЕНДЕРА', 0)
    doc.add_paragraph(f'Анализ проведён с использованием {prompt_info} промпта.')
    doc.add_paragraph(f'Номер тендера: {regNumber}')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
    doc.add_paragraph('=' * 50)

    doc.add_heading('Структурированный анализ', level=1)
    if isinstance(analysis_result, dict):
        for k, v in analysis_result.items():
            doc.add_paragraph(f'{k}: {v}')
    else:
        doc.add_paragraph(str(analysis_result))

    doc.add_page_break()
    doc.add_heading('КРИТИЧЕСКИЙ АНАЛИЗ И РЕКОМЕНДАЦИИ', level=1)
    doc.add_paragraph(critical_result if critical_result else "Не удалось выполнить критический анализ.")

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zf:
        zf.writestr(f'анализ_тендера_{regNumber}.docx', word_buffer.getvalue())
        for fname, content in original_files:
            zf.writestr(fname, content)

    zip_buffer.seek(0)
    filename = f'анализ_тендера_{regNumber}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip'
    encoded = quote(filename)

    logger.info(f"✅ [DONE] Тендер {regNumber} полностью обработан за {time.time() - start_total:.2f} сек")
    return Response(
        zip_buffer.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"}
    )

# ================= УПАКОВКА ФАЙЛОВ =================
@app.post("/package_files")
@limiter.limit("10/minute")
async def package_files(request: Request, files: list[UploadFile] = File(...), analysis_text: str = Form("")):
    await check_access(request)
    if not files:
        raise HTTPException(400, "Нет файлов")
    if len(files) > 100:
        raise HTTPException(400, "Слишком много файлов (макс. 100)")
    tenders = {}
    total_size = 0
    max_total_size = 500 * 1024 * 1024
    for file in files:
        content = await file.read()
        total_size += len(content)
        if total_size > max_total_size:
            raise HTTPException(400, "Общий размер файлов превышает 500 МБ")
        parts = file.filename.split('_', 1)
        tender_id = parts[0] if len(parts) == 2 else "без_тендера"
        original_name = parts[1] if len(parts) == 2 else file.filename
        file_type = detect_file_type(content, original_name)
        base = os.path.splitext(original_name)[0] or 'file'
        ext_map = {
            'pdf': '.pdf', 'xlsx': '.xlsx', 'xls': '.xls', 'docx': '.docx',
            'rar': '.rar', 'zip': '.zip', '7z': '.7z', 'png': '.png',
            'jpg': '.jpg', 'rtf': '.rtf', 'doc': '.doc'
        }
        new_ext = ext_map.get(file_type, '')
        if new_ext:
            original_name = base + new_ext
        else:
            if '.' not in original_name:
                original_name = base + '.bin'
        logger.info(f"Тип: {file_type}, имя: {original_name}")
        tenders.setdefault(tender_id, []).append((original_name, content))
    for tender_id, file_list in tenders.items():
        seen = set()
        new_list = []
        for name, content in file_list:
            base, ext = os.path.splitext(name)
            counter = 1
            new_name = name
            while new_name in seen:
                new_name = f"{base}_{counter}{ext}"
                counter += 1
            seen.add(new_name)
            new_list.append((new_name, content))
        tenders[tender_id] = new_list
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zf:
        for tender_id, file_list in tenders.items():
            folder = f"Тендер_{tender_id}"
            for name, content in file_list:
                zf.writestr(os.path.join(folder, name), content)
        if analysis_text:
            doc = Document()
            doc.add_heading('РЕЗУЛЬТАТЫ АНАЛИЗА ТЕНДЕРОВ', 0)
            doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
            doc.add_paragraph('=' * 50)
            sections = analysis_text.split('=== Тендер')
            for s in sections:
                if s.strip():
                    doc.add_paragraph(s.strip())
            word_buf = io.BytesIO()
            doc.save(word_buf)
            word_buf.seek(0)
            zf.writestr("анализ_тендеров.docx", word_buf.getvalue())
    zip_buffer.seek(0)
    filename = f"результаты_анализа_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    encoded = quote(filename)
    return Response(zip_buffer.getvalue(), media_type="application/zip", headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"})

# ================= ЭНДПОЙНТ ДЛЯ АНАЛИЗА ТЕКСТА ДЛЯ ГАРАНТИИ =================
@app.post("/api/guarantee/analyze")
@limiter.limit("10/minute")
async def guarantee_analyze(request: Request, data: dict):
    await check_access(request)
    reg_number = data.get("regNumber")
    text = data.get("text", "")
    selected_fields = data.get("fields", [])
    if not reg_number or not text or len(text) < 100:
        raise HTTPException(400, "Недостаточно данных для анализа")
    if not selected_fields:
        selected_fields = [
            "НАЗВАНИЕ АУКЦИОНА", "Начальная цена (НМЦ)", "ДОПОЛНИТЕЛЬНЫЕ ТРЕБОВАНИЯ К УЧАСТНИКУ",
            "ДАТА ОКОНЧАНИЯ/ПРОВЕДЕНИЯ", "Аванс", "Обеспечение заявки", "Обеспечение контракта",
            "Обеспечение гарантийных обязательств", "Контакты", "Место исполнения", "ДАТА ОКОНЧАНИЯ КОНТРАКТА"
        ]
    analysis_result = analyze_tender_text(text, selected_fields)
    await database.save_analysis_cache(reg_number, analysis_result)
    logger.info(f"Анализ для {reg_number} выполнен и сохранён в кэш")
    return {"status": "ok", "reg_number": reg_number}

# ================= ЭНДПОЙНТ ДЛЯ СТРАНИЦЫ ГАРАНТИИ =================
@app.get("/guarantee", response_class=HTMLResponse)
async def guarantee_page(request: Request):
    reg_number = request.query_params.get("regNumber")
    if not reg_number:
        return HTMLResponse("<h1>Ошибка</h1><p>Не указан номер тендера.</p>")
    if not reg_number.replace('-', '').replace('/', '').isalnum():
        raise HTTPException(400, "Некорректный номер тендера")
    cached = await database.get_cached_analysis(reg_number)
    data = cached if cached else {}

    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Заявка на банковскую гарантию</title>
        <style>
            body {{ font-family: Arial; padding: 20px; max-width: 600px; margin: auto; }}
            label {{ display: block; margin-top: 8px; font-weight: bold; font-size: 14px; }}
            input, select {{ width: 100%; padding: 6px; margin-top: 2px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px; box-sizing: border-box; }}
            .btn {{ background: #f59e0b; color: white; border: none; padding: 10px; font-size: 16px; border-radius: 4px; cursor: pointer; width: 100%; margin-top: 16px; }}
            .btn:hover {{ background: #d97706; }}
            .btn:disabled {{ background: #ccc; cursor: not-allowed; }}
            .field {{ margin-bottom: 12px; }}
            .inline-label {{
                display: flex;
                align-items: center;
                gap: 8px;
                margin: 6px 0 12px 0;
                justify-content: flex-start;
                font-weight: normal;
            }}
            .inline-label input[type="checkbox"] {{
                width: 16px;
                height: 16px;
                margin: 0;
                flex-shrink: 0;
            }}
            .inline-label span {{
                font-size: 13px;
                color: #1e293b;
            }}
            .consent-block {{
                margin-top: 10px;
                border-top: 1px solid #e2e8f0;
                padding-top: 12px;
            }}
            .consent-block label {{
                font-weight: normal;
                font-size: 13px;
                display: flex;
                align-items: flex-start;
                gap: 8px;
                margin-top: 6px;
            }}
            .consent-block input[type="checkbox"] {{
                width: 16px;
                height: 16px;
                margin-top: 2px;
                flex-shrink: 0;
            }}
        </style>
    </head>
    <body>
        <h1>Заявка на банковскую гарантию</h1>
        <form id="guaranteeForm" action="/api/guarantee/request" method="post">
            <input type="hidden" name="regNumber" value="{reg_number}">
            <div class="field">
                <label>Номер тендера</label>
                <input type="text" value="{reg_number}" readonly>
            </div>
            <div class="field">
                <label>Начальная цена (НМЦ)</label>
                <input type="text" name="nmc" value="{data.get('Начальная цена (НМЦ)', '')}" placeholder="Не указано">
            </div>
            <div class="field">
                <label>Дата окончания контракта</label>
                <input type="text" name="endDate" value="{data.get('ДАТА ОКОНЧАНИЯ КОНТРАКТА', '')}" placeholder="Не указано">
            </div>
            <div class="field">
                <label>Дата окончания подачи заявок</label>
                <input type="text" name="bidEndDate" value="{data.get('ДАТА ОКОНЧАНИЯ/ПРОВЕДЕНИЯ', '')}" placeholder="Не указано">
            </div>
            <div class="field">
                <label>Обеспечение заявки</label>
                <input type="text" name="bidSecurity" value="{data.get('Обеспечение заявки', '')}" placeholder="Не указано">
            </div>
            <div class="field">
                <label>Обеспечение контракта</label>
                <input type="text" name="contractSecurity" value="{data.get('Обеспечение контракта', '')}" placeholder="Не указано">
            </div>
            <div class="field">
                <label>Тип гарантии</label>
                <select name="guaranteeType" required>
                    <option value="">Выберите</option>
                    <option value="participation">Обеспечение заявки (участие)</option>
                    <option value="execution">Обеспечение исполнения контракта</option>
                </select>
            </div>
            <div class="field">
                <label>Ваше имя</label>
                <input type="text" name="clientName" placeholder="Иванов Иван Иванович" required>
            </div>
            <div class="field">
                <label>ИНН компании</label>
                <input type="text" name="inn" placeholder="1234567890" required pattern="[0-9]{{10,12}}">
            </div>
            <div class="field">
                <label>Телефон</label>
                <input type="tel" name="phone" placeholder="+7 (999) 123-45-67" required>
            </div>
            <div class="field">
                <label>Email</label>
                <input type="email" name="email" placeholder="user@example.com" required>
            </div>
            <div class="inline-label">
                <input type="checkbox" name="contact_by_email" value="true">
                <span>Не звонить мне, связываться только по email</span>
            </div>
            <div class="consent-block">
                <label>
                    <input type="checkbox" id="consent_personal" name="consent_personal" required>
                    Я даю согласие на обработку моих персональных данных в соответствии с Федеральным законом от 27.07.2006 № 152-ФЗ «О персональных данных»
                </label>
                <label>
                    <input type="checkbox" id="consent_terms" name="consent_terms" required>
                    Я принимаю условия <a href="https://csb24-tender.ru/offer" target="_blank">Пользовательского соглашения</a> и <a href="https://csb24-tender.ru/privacy-policy" target="_blank">Политики конфиденциальности</a>
                </label>
            </div>
            <button type="submit" class="btn" id="submitBtn" disabled>Отправить заявку</button>
        </form>
        <div style="margin-top: 20px; font-size: 13px; color: #666;">
            <a href="/">На главную</a>
        </div>
        <script>
            const consentPersonal = document.getElementById('consent_personal');
            const consentTerms = document.getElementById('consent_terms');
            const submitBtn = document.getElementById('submitBtn');
            function checkConsents() {{
                if (consentPersonal.checked && consentTerms.checked) {{
                    submitBtn.disabled = false;
                }} else {{
                    submitBtn.disabled = true;
                }}
            }}
            consentPersonal.addEventListener('change', checkConsents);
            consentTerms.addEventListener('change', checkConsents);
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html)

# ================= ЭНДПОЙНТ ПРИЁМА ЗАЯВКИ =================
@app.post("/api/guarantee/request")
@limiter.limit("5/minute")
async def guarantee_request(
    request: Request,
    regNumber: str = Form(...),
    nmc: str = Form(""),
    endDate: str = Form(""),
    bidEndDate: str = Form(""),
    bidSecurity: str = Form(""),
    contractSecurity: str = Form(""),
    guaranteeType: str = Form(...),
    clientName: str = Form(...),
    inn: str = Form(...),
    phone: str = Form(...),
    email: str = Form(...),
    contact_by_email: bool = Form(False),
    consent_personal: bool = Form(False),
    consent_terms: bool = Form(False)
):
    if not consent_personal or not consent_terms:
        raise HTTPException(400, "Необходимо дать согласие на обработку персональных данных и принять условия")
    inn_clean = inn.replace(' ', '').replace('-', '')
    if not inn_clean.isdigit() or len(inn_clean) not in [10, 12]:
        raise HTTPException(400, "Некорректный ИНН")
    if '@' not in email or '.' not in email.split('@')[1]:
        raise HTTPException(400, "Некорректный email")

    try:
        pool = await database.get_pool()
        async with pool.acquire() as conn:
            await conn.execute("""
                INSERT INTO guarantee_requests
                (reg_number, nmc, end_date, bid_end_date, guarantee_type,
                 client_name, inn, phone, email, bid_security, contract_security, contact_by_email)
                VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12)
            """, regNumber, nmc, endDate, bidEndDate, guaranteeType,
                clientName, inn, phone, email, bidSecurity, contractSecurity, contact_by_email)
        logger.info(f"Заявка для тендера {regNumber} сохранена в БД")

        asyncio.create_task(
            send_max_notification(
                reg_number=regNumber,
                client_name=clientName,
                inn=inn,
                phone=phone,
                email=email,
                nmc=nmc,
                end_date=endDate,
                bid_end_date=bidEndDate,
                bid_security=bidSecurity,
                contract_security=contractSecurity,
                guarantee_type=guaranteeType,
                contact_by_email=contact_by_email
            )
        )

    except Exception as e:
        logger.error(f"Ошибка сохранения в БД: {e}")
        raise HTTPException(500, "Ошибка при сохранении заявки")

    return HTMLResponse("""
    <!DOCTYPE html>
    <html>
    <head><meta charset="UTF-8"><title>Заявка отправлена</title></head>
    <body style="text-align: center; padding: 40px; font-family: Arial;">
        <h1 style="color: #10b981;">Заявка успешно отправлена!</h1>
        <p>Мы свяжемся с вами в ближайшее время.</p>
        <p style="margin-top: 20px;">
            <a href="/" style="color: #667eea; text-decoration: none;">На главную</a>
        </p>
    </body>
    </html>
    """)

# ================= ПРОБНЫЙ ПЕРИОД =================
@app.post("/api/trial/start")
@limiter.limit("3/hour")
async def start_trial(request: Request, data: TrialRequest):
    device_id = data.device_id
    ip_address = request.client.host
    user_agent = request.headers.get("User-Agent")
    result = await database.start_trial(device_id, trial_days=2, ip_address=ip_address, user_agent=user_agent)
    return result

@app.post("/api/trial/status")
@limiter.limit("10/minute")
async def check_trial(request: Request, data: TrialRequest):
    result = await database.get_trial_status(data.device_id)
    return result

# ================= ОСТАЛЬНЫЕ ЭНДПОЙНТЫ =================
@app.post("/search_tenders")
@limiter.limit("5/minute")
async def search_tenders(request: Request, data: dict):
    await check_access(request)
    query = data.get("query", "").strip()
    limit = data.get("limit", MAX_TENDERS)
    if not query:
        raise HTTPException(400, "Введите ключевые слова для поиска")
    if len(query) > 200:
        raise HTTPException(400, "Запрос слишком длинный")
    tender_urls = parser.search_tenders_zakupki(query, limit)
    if not tender_urls:
        return {"detail": "Тендеры по вашему запросу не найдены"}
    base_dir = f"search_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    os.makedirs(base_dir, exist_ok=True)
    results = []
    try:
        for idx, tender_url in enumerate(tender_urls[:MAX_TENDERS], 1):
            tender_name = f"Тендер_{idx}"
            tender_dir = os.path.join(base_dir, tender_name)
            os.makedirs(tender_dir, exist_ok=True)
            files = parser.download_files_from_tender(tender_url, tender_dir)
            if files:
                combined_text = ""
                for file_path in files:
                    text = read_docx(file_path) if file_path.endswith('.docx') else read_txt(file_path) if file_path.endswith('.txt') else read_excel(file_path)
                    if text and not text.startswith("Ошибка"):
                        combined_text += text + "\n"
                results.append({
                    "tender_name": tender_name,
                    "files": files,
                    "text": combined_text[:5000]
                })
            await asyncio.sleep(2)
    finally:
        if os.path.exists(base_dir):
            shutil.rmtree(base_dir, ignore_errors=True)
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zf:
        for res in results:
            doc = Document()
            doc.add_heading(f'Анализ тендера: {res["tender_name"]}', 0)
            doc.add_paragraph(f'Дата анализа: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
            doc.add_paragraph('=' * 50)
            doc.add_paragraph(res["text"] if res["text"] else "Не удалось извлечь текст из документов")
            word_buf = io.BytesIO()
            doc.save(word_buf)
            word_buf.seek(0)
            zf.writestr(f"{res['tender_name']}_результат.docx", word_buf.getvalue())
    zip_buffer.seek(0)
    filename = f"тендеры_по_запросу_{query[:20]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    encoded = quote(filename)
    return Response(zip_buffer.getvalue(), media_type="application/zip", headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"})

@app.post("/suggest_keywords")
@limiter.limit("10/minute")
async def suggest_keywords(request: Request, data: dict):
    await check_access(request)
    description = data.get("description", "").strip()
    if not description:
        raise HTTPException(400, "Описание не может быть пустым")
    if len(description) > 1000:
        raise HTTPException(400, "Описание слишком длинное")
    prompt = f"""Ты — помощник по тендерам. Пользователь описал свою деятельность:
"{description}"
Выдели 5–7 ключевых слов для поиска тендеров на zakupki.gov.ru.
Ключевые слова должны быть конкретными (например, "строительство школы", "поставка медоборудования", "ремонт дорог").
Выдай ТОЛЬКО список слов через запятую, без лишнего текста."""
    answer = query_deepseek(prompt)
    keywords = [kw.strip() for kw in answer.replace('\n', ',').split(',') if kw.strip()]
    return {"keywords": keywords[:7]}

@app.post("/ask_ai")
@limiter.limit("10/minute")
async def ask_ai(request: Request, data: dict):
    await check_access(request)
    question = data.get("question", "").strip()
    if not question:
        raise HTTPException(400, "Введите вопрос")
    if len(question) > 1000:
        raise HTTPException(400, "Вопрос слишком длинный")
    prompt = f"""Ты — консультант по тендерам и бизнес-процессам. Ответь на вопрос пользователя чётко, по делу, без воды.
Вопрос пользователя: {question}
Дай ответ в виде текста (2-4 предложения), который будет полезен для бизнесу."""
    answer = query_deepseek(prompt)
    return {"answer": answer}

# ================= ЭНДПОЙНТЫ ДЛЯ ЛИЦЕНЗИЙ =================
@app.post("/api/verify-license")
@limiter.limit("10/minute")
async def verify_license_endpoint(request: Request):
    license_key = request.headers.get("X-License-Key")
    if not license_key:
        return {"valid": False, "reason": "No license key provided"}

    result = await database.verify_license(license_key)
    if result and result.get("valid"):
        return {"valid": True}
    return {"valid": False, "reason": "Invalid or expired"}

@app.post("/api/activate-license")
@limiter.limit("5/minute")
async def activate_license(request: Request, data: LicenseActivateRequest):
    result = await database.verify_and_activate_license(data.key)
    if not result["valid"]:
        logger.warning(f"Неудачная активация лицензии: {data.key[:8]}... от {request.client.host}")
        return {"valid": False, "reason": result["reason"]}
    logger.info(f"Лицензия активирована: {data.key[:8]}...")
    return {"valid": True, "expires_at": result["expires_at"]}

# ================= ЗАПУСК =================
if __name__ == "__main__":
    port = int(os.getenv("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port)
